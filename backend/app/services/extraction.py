import os
from pathlib import Path


class ScannedPDFError(Exception):
    """
    Raised when a PDF has no extractable text layer anywhere in the document.

    Why this is a named exception and not a silent empty-string return:
    The caller — the background worker — needs to distinguish between
    "extraction failed due to a bug" and "this is a scanned image."
    They require different DB status values and different user-facing messages.
    A generic empty string conflates both cases and makes debugging impossible.

    Architecture note — why detection happens in the background worker, not
    the upload route:

    Option A (pre-flight in upload route): parse the file synchronously before
    returning 201. Problems: (1) even wrapped in asyncio.to_thread, you still
    parse the full document twice — once to check, once to extract in the
    background worker, doubling I/O on every upload. (2) A single-page heuristic
    to avoid the double-parse creates false rejections on PDFs with cover images.

    Option B (defer to background worker): extract once, raise ScannedPDFError
    if no text found anywhere, set status="failed_unreadable" in MongoDB.
    The frontend polls /documents/{id}/status and shows the error when it lands.
    Single I/O pass. No false rejections. User gets a 201 immediately, then a
    clear error message within seconds. This is the correct trade-off.
    """
    pass


def extract_text(file_path: str, file_type: str) -> str:
    """
    Extract raw text from PDF, DOCX, or TXT.
    Raises ScannedPDFError for image-only PDFs.
    Raises ValueError for unsupported types.
    Does NOT silently return empty string — callers must handle errors explicitly.
    """
    if file_type == "pdf":
        return _extract_pdf(file_path)
    elif file_type == "docx":
        return _extract_docx(file_path)
    elif file_type == "txt":
        return _extract_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def _extract_pdf(file_path: str) -> str:
    import fitz  # PyMuPDF
    text = []
    with fitz.open(file_path) as doc:
        for page in doc:
            page_text = page.get_text()
            if page_text.strip():
                text.append(page_text)

    result = "\n".join(text)
    if not result.strip():
        raise ScannedPDFError(
            "This PDF appears to be a scanned image with no extractable text. "
            "SmartDocs AI requires PDFs with a text layer. "
            "Try exporting from Word/Google Docs, or use an OCR tool first to create a searchable PDF."
        )
    return result


def _extract_docx(file_path: str) -> str:
    from docx import Document
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def _extract_txt(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="latin-1") as f:
            return f.read()


def get_file_type(filename: str) -> str:
    ext = Path(filename).suffix.lower().lstrip(".")
    allowed = {"pdf", "txt", "docx"}
    if ext not in allowed:
        raise ValueError(f"File type '{ext}' not supported. Allowed: {allowed}")
    return ext
